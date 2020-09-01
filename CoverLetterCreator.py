import os
import docx
import re
os.chdir('YOUR FILE PATH')

OpenFile = open('CoverLetterOutline.txt')

ListForm = OpenFile.readlines()
OpenFile.close()

print('''What's the company name?''')
CompanyName = input()

print('''What's the company's zipcode?''')
ZipCode = input()

print('''What's todays date?''')
print('''For example, write "May 31, 2020" instead of 6/31/2020''')
TodayDate = input()

print('Here is a list of paragraph options \n')
for i in range(1,len(ListForm)-7,2):
    print('Here is paragraph option number ' + str((i+1)/2) + ':')
    print(ListForm[i])

print('Choose the order you want your letter to be written')
print('For example, type 1234, 5342, 2341, etc... as the order you want them to be written:')
DesiredInput = input()

print('Great! We\'ll write this into a .docx file')

COMPANY_SUB = re.compile(r'COMPANY')
MATCH = COMPANY_SUB.sub(str(CompanyName), ListForm[14])
     # ListForm[14] has the expression where we want to make a subsitution

CreatingDoc = docx.Document('1 Cover Letter Prelude.docx')
     # We are including the prelude part of the cover letter
     # This is what we will write on in the following lines

# The following forloops replace statements with appropriate information
for paragraph in CreatingDoc.paragraphs:
    if 'COMPANY' in paragraph.text:
        paragraph.text = CompanyName

for paragraph in CreatingDoc.paragraphs:
    if 'ZIPCODE' in paragraph.text:
        paragraph.text = 'San Diego, CA ' + ZipCode

for paragraph in CreatingDoc.paragraphs:
    if 'DATE' in paragraph.text:
        paragraph.text = TodayDate

CreatingDoc.add_paragraph(ListForm[13])
CreatingDoc.add_paragraph(MATCH)
     # We are adding the beginning section of the cover letter

for i in range(len(DesiredInput)):
    CreatingDoc.add_paragraph(ListForm[((int(DesiredInput[i])*2)-1)])

CreatingDoc.add_paragraph(ListForm[9])
CreatingDoc.add_paragraph(ListForm[10])
CreatingDoc.add_paragraph(ListForm[11])
     # We are adding the ending section of the cover letter

CreatingDoc.save('Cover Letter ' + CompanyName + '.docx')    
print('Cover letter is created!')
    




